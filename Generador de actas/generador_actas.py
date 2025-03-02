import mysql.connector
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk  # Para usar Combobox
import pandas as pd
import sys
import re
import os
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pymysql
from datetime import datetime
from docx.shared import Pt
from io import BytesIO
from docx import Document
# Función para obtener la ruta del recurso
def obtener_ruta_recurso(relativa):
    if getattr(sys, 'frozen', False):  # Si está empaquetado con PyInstaller
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relativa)

# Conexión a la base de datos
conexion = mysql.connector.connect(
    host="sql10.freesqldatabase.com",
    user="sql10763670",
    password="ZLADxTPaZh",
    database="sql10763670",  # Agrega el nombre de la base de datos
    port=3306
)

cursor = conexion.cursor()

ruta_permiso = obtener_ruta_recurso(r"recursos\PERMISOS EXAMEN - 2023.docx")

ruta_acta = obtener_ruta_recurso(r"recursos\ACTA DE EXAMEN.docx")


class TablaPermisos(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.master = parent
        self.title("Registro de Permisos de Examen")
        self.geometry("1200x600")
        # Frame para los filtros
        frame_filtros = tk.Frame(self)
        frame_filtros.pack(pady=10, fill='x')
        
        # Filtro por DNI
        tk.Label(frame_filtros, text="Filtrar por DNI:").pack(side='left', padx=5)
        self.filtro_dni = tk.StringVar()
        tk.Entry(frame_filtros, textvariable=self.filtro_dni, width=15).pack(side='left', padx=5)
        
        # Filtro por Nombre
        tk.Label(frame_filtros, text="Filtrar por Nombre:").pack(side='left', padx=5)
        self.filtro_nombre = tk.StringVar()
        tk.Entry(frame_filtros, textvariable=self.filtro_nombre, width=20).pack(side='left', padx=5)
        
        # Filtro por Materia
        tk.Label(frame_filtros, text="Filtrar por Materia:").pack(side='left', padx=5)
        self.filtro_materia = tk.StringVar()
        self.combo_materias = ttk.Combobox(frame_filtros, 
                                         textvariable=self.filtro_materia,
                                         width=25)
        self.combo_materias.pack(side='left', padx=5)
        
        # Botón de búsqueda
        btn_buscar = tk.Button(frame_filtros, text="Buscar", command=self.cargar_datos)
        btn_buscar.pack(side='left', padx=5)
        
        # Botón limpiar filtros
        btn_limpiar = tk.Button(frame_filtros, text="Limpiar", command=self.limpiar_filtros)
        btn_limpiar.pack(side='left', padx=5)
        
        # Cargar materias en el combo
        self.cargar_materias()
        
        # Configurar el Treeview
        self.tree = ttk.Treeview(
            self, 
            columns=('DNI', 'Nombre', 'Materia', 'Especialidad', 'Curso', 'Condición', 
                    'id_permiso', 'materia_id', 'curso_original'),
            show='headings'
        )

        # Configurar encabezados visibles
        self.tree.heading('DNI', text='DNI')
        self.tree.heading('Nombre', text='Nombre')
        self.tree.heading('Materia', text='Materia')
        self.tree.heading('Especialidad', text='Especialidad')
        self.tree.heading('Curso', text='Curso')
        self.tree.heading('Condición', text='Condición')

        # Ocultar columnas técnicas
        for col in ['id_permiso', 'materia_id', 'curso_original']:
            self.tree.column(col, width=0, stretch=False)

        # Scrollbar
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # Layout
        self.tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Menú contextual
        self.menu = tk.Menu(self, tearoff=0)
        self.menu.add_command(label="Eliminar", command=self.eliminar_registro)
        self.menu.add_command(label="Modificar", command=self.modificar_registro)
        self.tree.bind("<Button-3>", self.mostrar_menu)       

        self.cargar_datos() 

    def cargar_materias(self):
        cursor.execute("SELECT DISTINCT nombre FROM materia")
        materias = [row[0] for row in cursor.fetchall()]
        self.combo_materias['values'] = materias

    def limpiar_filtros(self):
        self.filtro_dni.set('')
        self.filtro_nombre.set('')
        self.filtro_materia.set('')
        self.cargar_datos()


    def mostrar_menu(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.menu.post(event.x_root, event.y_root)
    
    def eliminar_registro(self):
        selected = self.tree.selection()
        if not selected:
            return
            
        item = self.tree.item(selected[0])
        valores = item['values']
        
        try:
            cursor.execute('''
                DELETE FROM detalle_permiso
                WHERE id_permiso = %s 
                AND materia = %s 
            ''', (valores[6], valores[7]))  # Índices correctos
            
            conexion.commit()
            self.cargar_datos()
            messagebox.showinfo("Éxito", "Registro eliminado correctamente")
            
        except mysql.connector.Error as err:
            conexion.rollback()
            messagebox.showerror("Error", f"No se pudo eliminar: {err}")
    
    def modificar_registro(self):
        selected = self.tree.selection()
        if not selected:
            return
            
        item = self.tree.item(selected[0])
        valores = item['values']
        
        # Ventana de modificación
        edit_win = tk.Toplevel(self)
        edit_win.title("Modificar Registro")
        
        # Variables del formulario
        dni = tk.StringVar(value=valores[0])
        nombre = tk.StringVar(value=valores[1])
        materia = tk.StringVar(value=f"{valores[7]} - {valores[2]}")  # ID - Nombre
        modalidad = tk.StringVar(value=valores[3])
        curso = tk.StringVar(value=valores[4])
        condicion = tk.StringVar(value=valores[5])

        # Función para actualizar materias según especialidad
        def actualizar_materias_combo(event=None):
            cursor.execute("SELECT id, nombre FROM materia WHERE modalidad = %s", 
                        (modalidad.get(),))
            materias = [f"{row[0]} - {row[1]}" for row in cursor.fetchall()]
            combo_materia['values'] = materias

        # Campos editables
        ttk.Label(edit_win, text="DNI:").grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(edit_win, text=valores[0]).grid(row=0, column=1, padx=5, pady=5)  # Display DNI as a label
        ttk.Label(edit_win, text="Nombre:").grid(row=1, column=0, padx=5, pady=5)
        ttk.Label(edit_win, text=valores[1]).grid(row=1, column=1, padx=5, pady=5) # Display Nombre as a label
        
        #Especialidad
        ttk.Label(edit_win, text="Especialidad:").grid(row=2, column=0, padx=5, pady=5)
        ttk.Label(edit_win, text=valores[3]).grid(row=2, column=1, padx=5, pady=5) # Display Nombre as a label

        #combo_modalidad = ttk.Combobox(edit_win, textvariable=modalidad, 
        #                           values=["MMO", "TEM", "POLIMODAL"], state='readonly')
        #combo_modalidad.grid(row=2, column=1, padx=5, pady=5)
        #combo_modalidad.bind("<<ComboboxSelected>>", actualizar_materias_combo)
        
        # Materia (dependiente de especialidad)
        ttk.Label(edit_win, text="Materia:").grid(row=3, column=0, padx=5, pady=5)
        combo_materia = ttk.Combobox(edit_win, textvariable=materia)
        combo_materia.grid(row=3, column=1, padx=5, pady=5)
        actualizar_materias_combo()  # Cargar inicialmente
        
        # Curso
        ttk.Label(edit_win, text="Curso:").grid(row=4, column=0, padx=5, pady=5)
        combo_curso = ttk.Combobox(edit_win, textvariable=curso, 
                                values=self.master.generar_cursos())
        combo_curso.grid(row=4, column=1, padx=5, pady=5)
        
        # Condición
        ttk.Label(edit_win, text="Condición:").grid(row=5, column=0, padx=5, pady=5)
        combo_condicion = ttk.Combobox(edit_win, textvariable=condicion, 
                                    values=["LIBRE", "REGULAR"])
        combo_condicion.grid(row=5, column=1, padx=5, pady=5)

        def guardar_cambios():
            try:
                # Extraer ID de materia
                materia_id = materia.get().split(" - ")[0]
                
                cursor.execute('''
                    UPDATE detalle_permiso 
                    SET condicion = %s, 
                        materia = %s,
                        curso = %s
                    WHERE id_permiso = %s 
                    AND materia = %s 
                ''', (condicion.get(), 
                    materia_id,
                    curso.get(),
                    valores[6],  # id_permiso
                    valores[7])) # materia_id original
                
                conexion.commit()               

                self.cargar_datos()
                edit_win.destroy()
                messagebox.showinfo("Éxito", "Registro actualizado correctamente")
                
            except mysql.connector.Error as err:
                conexion.rollback()
                messagebox.showerror("Error", f"Error al actualizar: {err}")
        
        ttk.Button(edit_win, text="Guardar Cambios", command=guardar_cambios).grid(row=6, columnspan=2, pady=10)

    def cargar_datos(self):
        try:
            # Construir consulta con filtros
            query = '''
                SELECT 
                    a.dni,
                    a.nombre,
                    m.nombre,
                    m.modalidad,
                    dp.curso,
                    dp.condicion,
                    p.nro,
                    m.id,
                    dp.curso
                FROM alumno a
                JOIN permiso p ON a.dni = p.dni
                JOIN detalle_permiso dp ON p.nro = dp.id_permiso
                JOIN materia m ON dp.materia = m.id
                WHERE 1=1
            '''
            
            params = []
            
            # Aplicar filtros
            if self.filtro_dni.get():
                query += " AND a.dni LIKE %s"
                params.append(f"%{self.filtro_dni.get()}%")
            
            if self.filtro_nombre.get():
                query += " AND a.nombre LIKE %s"
                params.append(f"%{self.filtro_nombre.get().upper()}%")
            
            if self.filtro_materia.get():
                query += " AND m.nombre = %s"
                params.append(self.filtro_materia.get())
            
            query += " ORDER BY a.nombre"
            
            cursor.execute(query, params)
            
            # Limpiar tabla existente
            for item in self.tree.get_children():
                self.tree.delete(item)
            
            # Insertar nuevos datos
            for row in cursor.fetchall():
                self.tree.insert('', 'end', values=row)
                
        except mysql.connector.Error as err:
            messagebox.showerror("Error", f"Error al cargar datos:\n{err}")

class RegistroMateriasApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Registro de Materias")
        self.geometry("1200x1200")
        
        self.dni = tk.StringVar()
        self.nombre = tk.StringVar()
        self.especialidad = tk.StringVar()
        self.curso = tk.StringVar() 
        self.materia = tk.StringVar()
        self.condicion = tk.StringVar()
        self.materias_seleccionadas = []
        
        self.crear_widgets()

        # Botón para mostrar tabla
        tk.Button(self, text="Ver Permisos Registrados", command=self.mostrar_tabla).pack(pady=10)
        # Botones en un Frame para alinearlos horizontalmente
        frame_botones = tk.Frame(self)
        frame_botones.pack(pady=10)

        tk.Button(frame_botones, text="Generar Permisos", command=self.generar_permiso).pack(side="left", padx=5)
        tk.Button(frame_botones, text="Generar Actas", command=self.generar_actas).pack(side="left", padx=5)
    
    def obtener_alumnos(self):
        cursor = conexion.cursor()
        # Consulta para obtener los DNI de los permisos
        query = "SELECT DISTINCT dni FROM permiso"
        cursor.execute(query)

        # Obtener los resultados en una lista
        dni_list = [row[0] for row in cursor.fetchall()]
        return dni_list

    def aplicar_estilo_encabezado(self, paragraph, texto, texto1, fuente="Calibri", tamaño=11, negrita=False):
        """
        Aplica el estilo Calibri de tamaño 11 solo a los encabezados, sin alterar el formato.
        """
        # Reemplazamos el texto sin cambiar el formato
        for run in paragraph.runs:
            if texto in run.text:  # Buscar el texto a reemplazar
                run.text = run.text.replace(run.text, texto1)
                run.font.name = fuente
                run.font.size = Pt(tamaño)
                run.font.bold = negrita

    def aplicar_estilo_personalizado_celda(self, celda, texto, fuente="Arial", tamaño=11, negrita=False):
        """
        Escribe texto en un párrafo dentro de una celda con estilo personalizado.
        """
        # Aseguramos que estamos trabajando con un párrafo en la celda
        p = celda.paragraphs[0]  # Primer párrafo en la celda
        p.clear()  # Limpiar el contenido previo
        run = p.add_run(texto)  # Agregar texto al párrafo
        run.font.name = fuente  # Tipo de letra
        run.font.size = Pt(tamaño)  # Tamaño de letra
        run.bold = negrita  # Negrita

        # Configuración explícita para garantizar el tipo de letra
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), fuente)
        rFonts.set(qn('w:hAnsi'), fuente)
        rFonts.set(qn('w:eastAsia'), fuente)
        rFonts.set(qn('w:cs'), fuente)
        rPr.append(rFonts)

    
    def generar_permiso_examen(self, id_alumno):
            cursor = conexion.cursor()
            # Obtener datos del alumno
            cursor.execute("""
                SELECT a.nombre, a.dni, p.nro, m.modalidad, m.nombre, dp.curso 
                FROM alumno a
                JOIN permiso p ON a.dni = p.dni
                JOIN detalle_permiso dp ON p.nro = dp.id_permiso
                JOIN materia m ON dp.materia = m.id
                WHERE a.dni = %s
            """, (id_alumno,))
            datos = cursor.fetchall()

            if not datos:
                return

            doc = Document(ruta_permiso)

            # Extraer datos del primer registro
            nombre_alumno, dni, nro_permiso, plan_estudio, materia, curso = datos[0]

            # Reemplazar valores en el documento
            # Llenar encabezados del acta
            for paragraph in doc.paragraphs:
                if "PERMISO DE EXAMEN N°" in paragraph.text:
                    self.aplicar_estilo_encabezado(paragraph, "56", str(nro_permiso), fuente="Arial", tamaño=12, negrita=True)
                if "alumno/a" in paragraph.text:
                    self.aplicar_estilo_encabezado(paragraph, "ORTIZ JOEL ALEXANDER", nombre_alumno, fuente="Arial", tamaño=12, negrita=True)
                if "DNI" in paragraph.text:
                    self.aplicar_estilo_encabezado(paragraph, "48663822", str(dni), fuente="Arial", tamaño=12, negrita=True)
                if "Plan de Estudios de" in paragraph.text:
                    self.aplicar_estilo_encabezado(paragraph, "TEM", plan_estudio, fuente="Arial", tamaño=12, negrita=True)

            # Llenar la tabla con las materias
            tabla = doc.tables[0]  # Primera tabla del documento
            orden = 1
            fila_actual = 1
            for i, (nombre_alumno, dni, nro_permiso, plan_estudio, materia, curso) in enumerate(datos, start=1):
                if pd.notna(materia):
                        if fila_actual < len(tabla.rows):
                            celda = tabla.rows[fila_actual].cells
                        else:
                            celda = tabla.add_row().cells

                        self.aplicar_estilo_personalizado_celda(celda[0], f"{orden:02}", fuente="Arial", tamaño=11, negrita=True)
                        self.aplicar_estilo_personalizado_celda(celda[1], materia.split(" (")[0], fuente="Arial", tamaño=11, negrita=True)
                        self.aplicar_estilo_personalizado_celda(celda[2], str(curso) if pd.notna(dni) else "", fuente="Arial",
                                                        tamaño=11, negrita=True)

                        for idx in [3, 4, 5]:
                            self.aplicar_estilo_personalizado_celda(celda[idx], "", fuente="Arial", tamaño=11, negrita=True)

                        orden += 1
                        fila_actual += 1
            for paragraph in doc.paragraphs:
                if "TINOGASTA, 30 de junio 		DE 2024" in paragraph.text:
                    fecha_actual = datetime.now().strftime("TINOGASTA, %d de %B DE %Y").capitalize()
                    paragraph.text = paragraph.text.replace("TINOGASTA, 30 de junio 		DE 2024", fecha_actual)
                    # Modificar el estilo de la fuente
                    run = paragraph.runs[0]  # El "run" representa un fragmento del texto en el párrafo
                    run.font.name = 'Arial'  # Cambiar la fuente a Arial
                    run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
                    run.font.bold = True  # Negrita
                

            # Guardar el documento modificado
            # Limpiar caracteres inválidos en el nombre del archivo
            alumno_limpio = re.sub(r'[<>:"/\\|?*]', '', nombre_alumno)

            # Crear la carpeta para la modalidad si no existe
            carpeta_alumno = "PERMISO_" + alumno_limpio.upper()
            if not os.path.exists(carpeta_alumno):
                os.makedirs(carpeta_alumno)

            # Generar el nombre del archivo
            nombre_archivo = f"Permiso_{alumno_limpio}.docx"

            # Generar la ruta completa
            ruta_archivo = os.path.join(carpeta_alumno, nombre_archivo)

            # Guardar el documento
            doc.save(ruta_archivo)
    
    def generar_permiso(self):
        try:
            dni_list = self.obtener_alumnos()
            for dni in dni_list:
                self.generar_permiso_examen(dni)
            messagebox.showinfo("Éxito", "Permisos generados exitosamente.")
        except Exception as e:
            messagebox.showerror("Error", e)

    def generar_actas_examen(self, modalidad):
        try:
            cursor = conexion.cursor(dictionary=True)
            
            # Consulta para obtener datos de alumnos y materias
            cursor.execute('''
                SELECT a.nombre AS alumno, a.dni, dp.curso, dp.condicion, m.nombre AS materia, m.modalidad FROM detalle_permiso dp 
                JOIN permiso p ON p.nro = dp.id_permiso 
                JOIN alumno a ON p.dni = a.dni
                JOIN materia m ON dp.materia = m.id
                WHERE m.modalidad = %s
            ''', (modalidad,))
            
            datos = cursor.fetchall()
            
            
            # Agrupar por materia, curso y condición
            grupos = {}
            for registro in datos:
                clave = (registro['materia'], registro['curso'], registro['condicion'])
                if clave not in grupos:
                    grupos[clave] = []
                grupos[clave].append(registro)
            # Procesar cada grupo
            for (materia, curso, condicion), alumnos in grupos.items():
                doc = Document(ruta_acta)
                
                # Llenar encabezados del acta
                for paragraph in doc.paragraphs:
                    if "EXAMEN DE ALUMNO" in paragraph.text:
                        self.aplicar_estilo_encabezado(paragraph, "REGULAR", condicion)
                    if "ESPACIO CURRICULAR" in paragraph.text:
                        self.aplicar_estilo_encabezado(paragraph, "HISTORIA III", materia.split(" (")[0])
                    if "PLAN DE ESTUDIO" in paragraph.text:
                        self.aplicar_estilo_encabezado(paragraph, "TEM", modalidad)
                    if "CURSO" in paragraph.text:
                        self.aplicar_estilo_encabezado(paragraph, "3º1º", curso)
                
                # Llenar tabla con alumnos
                tabla = doc.tables[0]
                fila_actual = 2  # Empezar desde la tercera fila (0-based)
                
                for i, alumno in enumerate(alumnos, start=1):
                    if fila_actual >= len(tabla.rows):
                        tabla.add_row()
                    
                    celdas = tabla.rows[fila_actual].cells
                    self.aplicar_estilo_personalizado_celda(celdas[0], str(i).zfill(2), fuente="Arial", tamaño=11, negrita=True)
                    self.aplicar_estilo_personalizado_celda(celdas[2], alumno['alumno'], fuente="Arial", tamaño=9, negrita=True)
                    self.aplicar_estilo_personalizado_celda(celdas[9], str(alumno['dni']) if pd.notna(alumno['dni']) else "", fuente="Arial", tamaño=9, negrita=True)

                    for idx in [1, 3, 4, 5, 6, 7, 8]:
                        self.aplicar_estilo_personalizado_celda(celdas[idx], "", fuente="Arial", tamaño=9, negrita=True)
                    
                    fila_actual += 1
                
                # Limpiar caracteres inválidos en el nombre del archivo
                materia_limpia = re.sub(r'[<>:"/\\|?*]', '', materia.split(" (")[0])
                curso_limpio = re.sub(r'[<>:"/\\|?*]', '', curso)
                modalidad_limpia = re.sub(r'[<>:"/\\|?*]', '', modalidad)
                condicion_limpia = re.sub(r'[<>:"/\\|?*]', '', condicion)

                # Crear la carpeta para la modalidad si no existe
                carpeta_modalidad = "ACTAS_"+modalidad.upper()
                if not os.path.exists(carpeta_modalidad):
                    os.makedirs(carpeta_modalidad)

                # Generar el nombre del archivo
                nombre_archivo = f"Acta_{modalidad_limpia}{materia_limpia}{curso_limpio}_{condicion_limpia}.docx"

                # Generar la ruta completa
                ruta_archivo = os.path.join(carpeta_modalidad, nombre_archivo)

                # Guardar el documento
                doc.save(ruta_archivo)
                
            messagebox.showinfo("Éxito", f"Actas {str(modalidad)} generadas correctamente")
        
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar actas: {str(e)}")
    
    def generar_actas(self):
        try:
            self.generar_actas_examen("TEM")
            self.generar_actas_examen("MMO")
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar actas: {str(e)}")

    def mostrar_tabla(self):
        TablaPermisos(self)

    def validate_dni(self, new_value):
        # Valida que sólo contenga números
        return new_value.isdigit() or new_value == ""

    
        
    def crear_widgets(self):
        # Campos para datos del alumno
        tk.Label(self, text="Nombre y Apellido:").pack(pady=5)
        tk.Entry(self, textvariable=self.nombre, width=40).pack(pady=5)
        
        tk.Label(self, text="DNI:").pack(pady=5)
        vcmd = (self.register(self.validate_dni), '%P')
        tk.Entry(self, textvariable=self.dni, validate='key', 
               validatecommand=vcmd, width=40).pack(pady=5)
        
        tk.Label(self, text="Especialidad:").pack(pady=5)
        self.combo_especialidad = ttk.Combobox(self, textvariable=self.especialidad, 
                                            values=["MMO", "TEM", "POLIMODAL"])
        self.combo_especialidad.pack(pady=5)
        self.combo_especialidad.bind("<<ComboboxSelected>>", self.actualizar_materias)
        
        # Combobox para materias
        tk.Label(self, text="Materia:").pack(pady=5)
        self.combo_materia = ttk.Combobox(self, textvariable=self.materia, width=40)
        self.combo_materia.pack(pady=5)

        # Combobox para curso (año + división)
        tk.Label(self, text="Curso:").pack(pady=5)
        self.combo_curso = ttk.Combobox(self, textvariable=self.curso, 
                                      values=self.generar_cursos())
        self.combo_curso.pack(pady=5)
        
        # Combobox para condición
        tk.Label(self, text="Condición:").pack(pady=5)
        self.combo_condicion = ttk.Combobox(self, textvariable=self.condicion, 
                                         values=["LIBRE", "REGULAR"])
        self.combo_condicion.pack(pady=5)

        # Botones en un Frame para alinearlos horizontalmente
        frame_botones_materias = tk.Frame(self)
        frame_botones_materias.pack(pady=10)

        tk.Button(frame_botones_materias, text="Agregar Materia", command=self.agregar_materia).pack(side="left", padx=5)
        tk.Button(frame_botones_materias, text="Eliminar Materia Seleccionada", command=self.eliminar_materia).pack(side="left", padx=5)

        
        # Listbox para materias agregadas
        self.listbox = tk.Listbox(self, width=50, height=10)
        self.listbox.pack(pady=10)
        
        # Botón para guardar
        tk.Button(self, text="Guardar Registro", command=self.guardar_registro).pack(pady=10)

    def eliminar_materia(self):
        selected_index = self.listbox.curselection()
        if selected_index:
            index = selected_index[0]
            del self.materias_seleccionadas[index]
            self.listbox.delete(index)
        else:
            messagebox.showinfo("Información", "Seleccione una materia para eliminar.")
    
    def generar_cursos(self):
        """Genera todas las combinaciones de año y división"""
        cursos = []
        for anio in range(1, 8):  # Años del 1 al 7
            for division in range(1, 3):  # Divisiones 1 y 2
                cursos.append(f"{anio}°{division}°")
        return cursos
    
    def actualizar_materias(self, event=None):
        especialidad = self.especialidad.get()
        cursor.execute("SELECT id, nombre FROM materia WHERE modalidad = %s", (especialidad,))
        materias = [f"{row[0]} - {row[1]}" for row in cursor.fetchall()]
        self.combo_materia['values'] = materias
    
    def agregar_materia(self):
        materia = self.combo_materia.get()
        condicion = self.condicion.get()
        if materia and condicion:
            self.materias_seleccionadas.append((materia.split(" - ")[0], condicion))
            self.listbox.insert(tk.END, f"{materia.split(' - ')[1]} ({condicion})")
    
    def validar_campos(self):
        errores = []

        self.cursor = conexion.cursor()
        self.cursor.execute("SELECT modalidad FROM alumno WHERE dni = %s", (self.dni.get(),))
        resultado = self.cursor.fetchone()
        
        # Validar campos básicos
        if not self.nombre.get().strip():
            errores.append("Nombre y Apellido")
        if not self.dni.get().strip():
            errores.append("DNI")
        if not self.especialidad.get():
            errores.append("Especialidad")
        if not self.materia.get():
            errores.append("Materia")
        if not self.condicion.get():
            errores.append("Condición")
        if not self.curso.get():
            errores.append("Curso")
        
        # Validar materias agregadas
        if not self.materias_seleccionadas:
            errores.append("Debe agregar al menos una materia")
        
        # Mostrar errores si existen
        if errores:
            mensaje = "Faltan completar los siguientes campos:\n"
            mensaje += "\n".join(f"- {campo}" for campo in errores)
            messagebox.showerror("Error de validación", mensaje)
            return False
        
        if not resultado:
            pass
        else:
            modalidad_bd = resultado[0]
            if modalidad_bd != self.especialidad.get():
                messagebox.showerror("Error", "La modalidad seleccionada no es la misma que que tiene registrada el alumno.")
                return False
        
        return True
    
    def guardar_registro(self):

        if not self.validar_campos():
            return
        # Validación adicional del DNI
        if not self.dni.get().isdigit() or len(self.dni.get()) < 7:
            messagebox.showerror("Error", "DNI inválido: Debe contener sólo números y tener al menos 7 dígitos")
            return

        try: 
            # Guardar alumno
            dni = self.dni.get()
            nombre = self.nombre.get()
            curso = self.curso.get()
            especialidad = self.curso.get()
        
            cursor.execute("INSERT IGNORE INTO alumno VALUES (%s, %s, %s)", (dni, nombre.upper(), especialidad.upper()))
        
            # 2. Obtener o crear permiso
            cursor.execute("SELECT nro FROM permiso WHERE dni = %s", (dni,))
            permiso_existente = cursor.fetchone()

            if permiso_existente:
                id_permiso = permiso_existente[0]
            else:
                cursor.execute("INSERT INTO permiso (dni) VALUES (%s)", (dni,))
                id_permiso = cursor.lastrowid
        
            # Guardar materias
            for materia_id, condicion in self.materias_seleccionadas:
                cursor.execute("INSERT INTO detalle_permiso VALUES (%s, %s, %s, %s)", 
                         (id_permiso, int(materia_id), condicion, curso))
            conexion.commit()
        
            # Limpiar formulario
            self.dni.set('')
            self.nombre.set('')
            self.especialidad.set('')
            self.materia.set('')
            self.condicion.set('')
            self.curso.set('')
            self.listbox.delete(0, tk.END)
            self.materias_seleccionadas = []
        
            tk.messagebox.showinfo("Éxito", "Registro guardado correctamente")
        except mysql.connector.Error as err:
            conexion.rollback()
            messagebox.showerror("Error", f"Error de base de datos:\n{err}") 

if __name__ == "__main__":
    app = RegistroMateriasApp()
    app.mainloop()
    conexion.close()