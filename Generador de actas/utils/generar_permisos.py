from tkinter import messagebox, filedialog
import pandas as pd
import re
import os
from docx.shared import Pt
from sqlalchemy import extract, and_
from sqlalchemy.exc import SQLAlchemyError
from modelos.init import Materia, DetallePermiso, Permiso, Alumno
from datetime import datetime
from docx import Document
from .utils import obtener_ruta_recurso, aplicar_estilo_encabezado, aplicar_estilo_personalizado_celda

ruta_permiso = obtener_ruta_recurso(r"recursos\PERMISOS EXAMEN - 2023.docx")

def generar_permiso_examen(session, id_alumno, base_path):
    try:
        current_date = datetime.now()
        mes_actual = current_date.month
        año_actual = current_date.year

        # Consulta ORM corregida
        datos = session.query(
            Alumno.nombre,
            Alumno.dni,
            Permiso.nro,
            Materia.modalidad,
            Materia.nombre,
            DetallePermiso.curso
        ).join(Permiso, Alumno.permisos)\
        .join(DetallePermiso, Permiso.detalles)\
        .join(Materia, DetallePermiso.materia_rel)\
        .filter(
            and_(
                Alumno.dni == id_alumno,
                extract('month', Permiso.fechaPermiso) == mes_actual,
                extract('year', Permiso.fechaPermiso) == año_actual
            )
        ).all()

        if not datos:
            return

        # Procesamiento del documento (igual que antes)
        doc = Document(ruta_permiso)
        nombre_alumno, dni, nro_permiso, plan_estudio, materia, curso = datos[0]

        # Reemplazar valores en el documento
        # Llenar encabezados del acta
        for paragraph in doc.paragraphs:
            if "PERMISO DE EXAMEN N°" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "56", str(nro_permiso), fuente="Arial", tamaño=12, negrita=True)
            if "alumno/a" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "ORTIZ JOEL ALEXANDER", nombre_alumno, fuente="Arial", tamaño=12, negrita=True)
            if "DNI" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "48663822", str(dni), fuente="Arial", tamaño=12, negrita=True)
            if "Plan de Estudios de" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "TEM", plan_estudio, fuente="Arial", tamaño=12, negrita=True)

        # Llenar la tabla con las materias
        tabla = doc.tables[0]  # Primera tabla del documento
        orden = 1
        fila_actual = 1
        for i, (nombre_alumno, dni, nro_permiso, plan_estudio, materia, curso) in enumerate(datos, start=1):
            if pd.notna(materia):
                    if fila_actual < len(tabla.rows):
                        celda = tabla.rows[fila_actual].cells
                    else:
                        celda = tabla.add_row().cells

                    aplicar_estilo_personalizado_celda(celda[0], f"{orden:02}", fuente="Arial", tamaño=11, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[1], materia.split(" (")[0], fuente="Arial", tamaño=11, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[2], str(curso) if pd.notna(dni) else "", fuente="Arial",
                                                    tamaño=11, negrita=True)

                    for idx in [3, 4, 5]:
                        aplicar_estilo_personalizado_celda(celda[idx], "", fuente="Arial", tamaño=11, negrita=True)

                    orden += 1
                    fila_actual += 1
        for paragraph in doc.paragraphs:
            if "TINOGASTA, 30 de junio 		DE 2024" in paragraph.text:
                fecha_actual = datetime.now().strftime("TINOGASTA, %d de %B DE %Y").capitalize()
                paragraph.text = paragraph.text.replace("TINOGASTA, 30 de junio 		DE 2024", fecha_actual)
                # Modificar el estilo de la fuente
                run = paragraph.runs[0]  # El "run" representa un fragmento del texto en el párrafo
                run.font.name = 'Arial'  # Cambiar la fuente a Arial
                run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
                run.font.bold = True  # Negrita
            

        # Guardar el documento modificado
        # Limpiar caracteres inválidos en el nombre del archivo
        alumno_limpio = re.sub(r'[<>:"/\\|?*]', '', nombre_alumno)

        # Crear la carpeta para la modalidad si no existe
        carpeta_alumno = os.path.join(base_path, f"PERMISO_{alumno_limpio.upper()}")
        if not os.path.exists(carpeta_alumno):
            os.makedirs(carpeta_alumno)

        # Generar el nombre del archivo
        nombre_archivo = f"Permiso_{alumno_limpio}.docx"

        # Generar la ruta completa
        ruta_archivo = os.path.join(carpeta_alumno, nombre_archivo)

        # Guardar el documento
        doc.save(ruta_archivo)
    except SQLAlchemyError as e:
        messagebox.showerror("Error", f"Error de base de datos: {str(e)}")
    except Exception as e:
        messagebox.showerror("Error", f"Error inesperado: {str(e)}")

def obtener_alumnos(session):
    try:
        # Consulta de DNIs únicos usando ORM
        return [resultado[0] for resultado in session.query(Permiso.dni).distinct().all()]
    except SQLAlchemyError as e:
        messagebox.showerror("Error", f"Error al obtener alumnos: {str(e)}")
        return []
    
def generar_permiso(session):
    try:
        base_path = filedialog.askdirectory(title="Seleccionar carpeta para guardar los permisos")
        if not base_path:
            messagebox.showwarning("Cancelado", "Operación cancelada por el usuario")
            return
            
        dni_list = obtener_alumnos(session)
        for dni in dni_list:
            generar_permiso_examen(session, dni, base_path)
            
        messagebox.showinfo("Éxito", f"✅ Permisos generados en: {base_path}")
        
    except Exception as e:
        messagebox.showerror("Error", f"Error al generar permisos: {str(e)}")