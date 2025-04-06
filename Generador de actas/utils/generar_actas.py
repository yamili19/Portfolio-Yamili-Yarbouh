from tkinter import messagebox
import pandas as pd
import re
import os
from docx import Document
from .utils import obtener_ruta_recurso, aplicar_estilo_encabezado, aplicar_estilo_personalizado_celda
from tkinter import filedialog  
from datetime import datetime
from modelos.init import Materia, DetallePermiso, Permiso, Alumno
from sqlalchemy import extract, and_
from sqlalchemy.exc import SQLAlchemyError


ruta_acta = obtener_ruta_recurso(r"recursos\ACTA DE EXAMEN.docx")

def generar_actas_examen(session, modalidad, base_path):
    try:
        current_date = datetime.now()
        mes_actual = current_date.month
        año_actual = current_date.year

        # Consulta ORM con joins y filtros
        query = session.query(
            Alumno.nombre.label('alumno'),
            Alumno.dni,
            DetallePermiso.curso,
            DetallePermiso.condicion,
            Materia.nombre.label('materia'),
            Materia.modalidad
        ).join(Permiso, DetallePermiso.permiso)\
         .join(Alumno, Permiso.alumno)\
         .join(Materia, DetallePermiso.materia_rel)\
         .filter(
             and_(
                 Materia.modalidad == modalidad,
                 extract('month', Permiso.fechaPermiso) == mes_actual,
                 extract('year', Permiso.fechaPermiso) == año_actual
             )
         )

        # Convertir resultados a formato de diccionario
        datos = [{
            'alumno': r.alumno,
            'dni': r.dni,
            'curso': r.curso,
            'condicion': r.condicion,
            'materia': r.materia,
            'modalidad': r.modalidad
        } for r in query.all()]

        # Agrupar por materia, curso y condición (igual que antes)
        grupos = {}
        for registro in datos:
            clave = (registro['materia'], registro['curso'], registro['condicion'])
            grupos.setdefault(clave, []).append(registro)

        # Procesar cada grupo
        for (materia, curso, condicion), alumnos in grupos.items():
            doc = Document(ruta_acta)
            
            # Llenar encabezados del acta
            for paragraph in doc.paragraphs:
                if "EXAMEN DE ALUMNO" in paragraph.text:
                    aplicar_estilo_encabezado(paragraph, "REGULAR", condicion)
                if "ESPACIO CURRICULAR" in paragraph.text:
                    aplicar_estilo_encabezado(paragraph, "HISTORIA III", materia.split(" (")[0])
                if "PLAN DE ESTUDIO" in paragraph.text:
                    aplicar_estilo_encabezado(paragraph, "TEM", modalidad)
                if "CURSO" in paragraph.text:
                    aplicar_estilo_encabezado(paragraph, "3º1º", curso)
            
            # Llenar tabla con alumnos
            tabla = doc.tables[0]
            fila_actual = 2  # Empezar desde la tercera fila (0-based)
            
            for i, alumno in enumerate(alumnos, start=1):
                if fila_actual >= len(tabla.rows):
                    tabla.add_row()
                
                celdas = tabla.rows[fila_actual].cells
                aplicar_estilo_personalizado_celda(celdas[0], str(i).zfill(2), fuente="Arial", tamaño=11, negrita=True)
                aplicar_estilo_personalizado_celda(celdas[2], alumno['alumno'], fuente="Arial", tamaño=9, negrita=True)
                aplicar_estilo_personalizado_celda(celdas[9], str(alumno['dni']) if pd.notna(alumno['dni']) else "", fuente="Arial", tamaño=9, negrita=True)

                for idx in [1, 3, 4, 5, 6, 7, 8]:
                    aplicar_estilo_personalizado_celda(celdas[idx], "", fuente="Arial", tamaño=9, negrita=True)
                
                fila_actual += 1
            
            # Limpiar caracteres inválidos en el nombre del archivo
            materia_limpia = re.sub(r'[<>:"/\\|?*]', '', materia.split(" (")[0])
            curso_limpio = re.sub(r'[<>:"/\\|?*]', '', curso)
            modalidad_limpia = re.sub(r'[<>:"/\\|?*]', '', modalidad)
            condicion_limpia = re.sub(r'[<>:"/\\|?*]', '', condicion)

            carpeta_modalidad = os.path.join(base_path, f"ACTAS_{modalidad.upper()}")
            if not os.path.exists(carpeta_modalidad):
                os.makedirs(carpeta_modalidad)

            # Generar el nombre del archivo
            nombre_archivo = f"Acta_{modalidad_limpia}_{materia_limpia}_{curso_limpio}_{condicion_limpia}.docx"

            # Generar la ruta completa
            ruta_archivo = os.path.join(carpeta_modalidad, nombre_archivo)

            # Guardar el documento
            doc.save(ruta_archivo)
                
        messagebox.showinfo("Éxito", f"✅ Actas {str(modalidad)} generadas y guardadas correctamente en {base_path}")
    except SQLAlchemyError as e:
        messagebox.showerror("Error", f"🚨 Error de base de datos: {str(e)}")
    except Exception as e:
        messagebox.showerror("Error", f"🚨 Error inesperado: {str(e)}")
        
    
def generar_actas(session):
    try:
        base_path = filedialog.askdirectory(
            title="Seleccionar carpeta para guardar las actas",
            mustexist=True
        )
        
        if not base_path:
            messagebox.showwarning("Cancelado", "Operación cancelada por el usuario")
            return
            
        generar_actas_examen(session, "TEM", base_path)
        generar_actas_examen(session, "MMO", base_path)
        
    except Exception as e:
        messagebox.showerror("Error", f"Error al generar actas: {str(e)}")